#!/usr/bin/env python3
"""
Advanced DOCX File Translator

This script automates the translation of Microsoft Word (.docx) files using the AsyncTranslator.
It preserves document formatting, styles, images, tables, and other elements while translating
text content from one language to another, with special handling for RTL languages.

Usage:
    python docx_translator.py -i input.docx -o output.docx -s en -t ar

Features:
- Preserves document formatting and structure
- Handles tables, lists, and nested elements
- Special handling for RTL languages (Arabic, Hebrew, etc.)
- Batch processing of multiple files
- Progress tracking and detailed reporting
- Translation memory to avoid redundant translations
"""

import os
import sys
import argparse
import asyncio
import logging
import time
import re
from pathlib import Path
from typing import Dict, List, Tuple, Set, Optional, Any, Union
from dataclasses import dataclass, field

import docx
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from translator import AsyncTranslator, TranslationResult  # Import the AsyncTranslator we created earlier

# RTL language codes
RTL_LANGUAGES = {'ar', 'he', 'fa', 'ur', 'ku', 'sd', 'ps', 'yi', 'dv'}

@dataclass
class TranslationStats:
    """Statistics about the translation process."""
    total_paragraphs: int = 0
    total_text_chunks: int = 0
    translated_chunks: int = 0
    skipped_chunks: int = 0
    failed_chunks: int = 0
    total_words: int = 0
    translated_words: int = 0
    time_taken: float = 0.0
    character_count: int = 0

@dataclass
class TranslationMemory:
    """Simple translation memory to avoid redundant translations."""
    entries: Dict[str, str] = field(default_factory=dict)
    hits: int = 0
    misses: int = 0
    
    def get(self, text: str) -> Optional[str]:
        """Get translation from memory if available."""
        if text in self.entries:
            self.hits += 1
            return self.entries[text]
        self.misses += 1
        return None
        
    def add(self, source: str, translation: str) -> None:
        """Add a translation to memory."""
        self.entries[source] = translation
        
    def save(self, filename: str) -> None:
        """Save translation memory to a file."""
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                for source, translation in self.entries.items():
                    f.write(f"{source}\t{translation}\n")
        except Exception as e:
            logging.error(f"Failed to save translation memory: {e}")
    
    @classmethod
    def load(cls, filename: str) -> 'TranslationMemory':
        """Load translation memory from a file."""
        memory = cls()
        try:
            if os.path.exists(filename):
                with open(filename, 'r', encoding='utf-8') as f:
                    for line in f:
                        parts = line.strip().split('\t', 1)
                        if len(parts) == 2:
                            memory.add(parts[0], parts[1])
        except Exception as e:
            logging.error(f"Failed to load translation memory: {e}")
        return memory

class DocxTranslator:
    """
    A class to translate DOCX files while preserving their structure and formatting.
    Uses AsyncTranslator for the actual translations.
    """
    
    def __init__(self, 
                 source_lang: str = 'en', 
                 target_lang: str = 'es',
                 max_concurrent: int = 10,
                 logger: Optional[logging.Logger] = None,
                 translation_memory_file: Optional[str] = None):
        """
        Initialize the DOCX translator.
        
        Args:
            source_lang: Source language code
            target_lang: Target language code
            max_concurrent: Maximum number of concurrent translation requests
            logger: Custom logger (if None, a default one will be created)
            translation_memory_file: Path to translation memory file
        """
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.max_concurrent = max_concurrent
        self.is_target_rtl = target_lang in RTL_LANGUAGES
        
        # Set up logger
        if logger is None:
            self.logger = logging.getLogger('docx_translator')
            self.logger.setLevel(logging.INFO)
            
            # Console handler
            console_handler = logging.StreamHandler()
            console_handler.setFormatter(logging.Formatter('%(levelname)s: %(message)s'))
            self.logger.addHandler(console_handler)
        else:
            self.logger = logger
            
        # Initialize the async translator
        self.translator = AsyncTranslator(
            max_concurrent_requests=max_concurrent,
            logger=self.logger
        )
        
        # Load translation memory if provided
        if translation_memory_file:
            self.memory = TranslationMemory.load(translation_memory_file)
            self.memory_file = translation_memory_file
        else:
            self.memory = TranslationMemory()
            self.memory_file = None
            
        # Minimum length for text to be considered for translation
        self.min_text_length = 1
        
        # Regex for finding words (used for word count)
        self.word_regex = re.compile(r'\b\w+\b')
    
    async def translate_docx(self, 
                           input_file: str, 
                           output_file: str,
                           batch_size: int = 50) -> TranslationStats:
        """
        Translate a DOCX file while preserving its structure.
        
        Args:
            input_file: Path to input DOCX file
            output_file: Path to output DOCX file
            batch_size: Size of translation batches
            
        Returns:
            TranslationStats object with statistics about the translation process
        """
        start_time = time.time()
        stats = TranslationStats()
        
        try:
            # Load the document
            self.logger.info(f"Loading DOCX file: {input_file}")
            doc = Document(input_file)
            
            # Extract all text content that needs translation
            self.logger.info("Extracting text content for translation")
            elements, text_chunks = self.extract_text_from_document(doc)
            
            # Count total words
            for text in text_chunks:
                stats.total_words += len(self.word_regex.findall(text))
                stats.character_count += len(text)
            
            # Update stats
            stats.total_paragraphs = len(elements)
            stats.total_text_chunks = len(text_chunks)
            
            if not text_chunks:
                self.logger.info("No text content found for translation. Saving document as is.")
                doc.save(output_file)
                stats.time_taken = time.time() - start_time
                return stats
            
            # Prepare chunks for translation (filter out those in translation memory)
            chunks_to_translate = []
            memory_translations = []
            
            for text in text_chunks:
                mem_trans = self.memory.get(text)
                if mem_trans:
                    memory_translations.append(mem_trans)
                    stats.skipped_chunks += 1
                else:
                    chunks_to_translate.append(text)
                    memory_translations.append(None)
            
            # If there are chunks to translate
            if chunks_to_translate:
                self.logger.info(f"Translating {len(chunks_to_translate)} text chunks "
                                f"from {self.source_lang} to {self.target_lang}")
                
                # Translate the texts
                translation_results = await self.translator.batch_translate_large(
                    texts=chunks_to_translate,
                    src_lang=self.source_lang,
                    dest_lang=self.target_lang,
                    batch_size=batch_size
                )
                
                # Count translated words
                for result in translation_results:
                    if result.success and result.translated_text:
                        stats.translated_chunks += 1
                        stats.translated_words += len(self.word_regex.findall(result.translated_text))
                    else:
                        stats.failed_chunks += 1
                
                # Merge results back with memory translations
                all_translations = []
                result_index = 0
                
                for mem_trans in memory_translations:
                    if mem_trans:
                        all_translations.append(mem_trans)
                    else:
                        result = translation_results[result_index]
                        result_index += 1
                        
                        if result.success and result.translated_text:
                            all_translations.append(result.translated_text)
                            # Add to translation memory
                            source_text = chunks_to_translate[result_index - 1]
                            self.memory.add(source_text, result.translated_text)
                        else:
                            # If translation failed, use original text
                            all_translations.append(chunks_to_translate[result_index - 1])
            else:
                # All translations were found in memory
                self.logger.info("All translations found in translation memory")
                all_translations = memory_translations
            
            # Replace original text with translations
            self.logger.info("Replacing text content with translations")
            self.replace_text_in_document(doc, elements, all_translations)
            
            # Apply RTL formatting if target language is RTL
            if self.is_target_rtl:
                self.logger.info("Applying RTL formatting")
                self.apply_rtl_formatting(doc)
            
            # Save the translated document
            self.logger.info(f"Saving translated document to: {output_file}")
            doc.save(output_file)
            
            # Save translation memory if needed
            if self.memory_file:
                self.memory.save(self.memory_file)
                self.logger.info(f"Translation memory updated (hits: {self.memory.hits}, misses: {self.memory.misses})")
            
            stats.time_taken = time.time() - start_time
            
            # Log summary
            self.logger.info(
                f"Translation completed in {stats.time_taken:.2f}s: "
                f"Translated {stats.translated_chunks}/{stats.total_text_chunks} chunks, "
                f"{stats.translated_words}/{stats.total_words} words"
            )
            
            return stats
            
        except Exception as e:
            self.logger.error(f"Error translating DOCX file: {str(e)}")
            stats.time_taken = time.time() - start_time
            raise
    
    def extract_text_from_document(self, doc: Document) -> Tuple[List[Any], List[str]]:
        """
        Extract all text content from a document that needs translation.
        
        Args:
            doc: The docx Document object
            
        Returns:
            A tuple of (elements, text_chunks) where elements is a list of elements
            that contain text and text_chunks is a list of text strings to translate
        """
        elements = []
        text_chunks = []
        
        # Process document body
        for element in self._iter_block_items(doc):
            if isinstance(element, Paragraph):
                text = element.text.strip()
                if text and len(text) > self.min_text_length:
                    elements.append(element)
                    text_chunks.append(text)
            elif isinstance(element, Table):
                for row in element.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text.strip()
                            if text and len(text) > self.min_text_length:
                                elements.append(paragraph)
                                text_chunks.append(text)
        
        # Process headers and footers
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                text = paragraph.text.strip()
                if text and len(text) > self.min_text_length:
                    elements.append(paragraph)
                    text_chunks.append(text)
                    
            # Footer
            for paragraph in section.footer.paragraphs:
                text = paragraph.text.strip()
                if text and len(text) > self.min_text_length:
                    elements.append(paragraph)
                    text_chunks.append(text)
        
        return elements, text_chunks
    
    def replace_text_in_document(self, doc: Document, elements: List[Any], translations: List[str]) -> None:
        """
        Replace the original text in the document with translations.
        
        Args:
            doc: The docx Document object
            elements: List of elements that contain text
            translations: List of translated text strings
        """
        for element, translation in zip(elements, translations):
            if isinstance(element, Paragraph):
                # Clear the paragraph
                for run in list(element.runs):
                    element._p.remove(run._r)
                
                # Add the translation as a new run
                element.add_run(translation)
    
    def apply_rtl_formatting(self, doc: Document) -> None:
        """
        Apply right-to-left formatting to the document.
        
        Args:
            doc: The docx Document object
        """
        # Set document direction for all paragraphs
        for paragraph in doc.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            paragraph.paragraph_format.right_indent = 0
            paragraph.paragraph_format.left_indent = 0
            
            # Set BiDi properties at the XML level
            p = paragraph._p
            p.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi", "1")
        
        # Handle paragraphs in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        p = paragraph._p
                        p.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi", "1")
    
    def _iter_block_items(self, doc: Document) -> List[Any]:
        """
        Iterate through all block items in the document.
        This includes paragraphs and tables.
        
        Args:
            doc: The docx Document object
            
        Returns:
            Generator that yields paragraphs and tables
        """
        for block in doc.element.body:
            if isinstance(block, CT_P):
                yield Paragraph(block, doc)
            elif isinstance(block, CT_Tbl):
                yield Table(block, doc)
    
    async def batch_translate_docx_files(self, 
                                       input_files: List[str], 
                                       output_dir: str) -> Dict[str, TranslationStats]:
        """
        Translate multiple DOCX files.
        
        Args:
            input_files: List of input DOCX file paths
            output_dir: Directory where translated files will be saved
            
        Returns:
            Dictionary mapping file paths to their TranslationStats
        """
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        
        all_stats = {}
        
        for input_file in input_files:
            filename = os.path.basename(input_file)
            output_file = os.path.join(output_dir, filename)
            
            self.logger.info(f"Processing file: {input_file} -> {output_file}")
            
            try:
                stats = await self.translate_docx(
                    input_file=input_file,
                    output_file=output_file
                )
                all_stats[input_file] = stats
            except Exception as e:
                self.logger.error(f"Failed to translate {input_file}: {str(e)}")
                all_stats[input_file] = TranslationStats(failed_chunks=1)
        
        return all_stats

async def main():
    """Main function."""
    parser = argparse.ArgumentParser(description='Translate DOCX files with RTL support')
    
    # Basic arguments
    parser.add_argument('-i', '--input', required=True, help='Input DOCX file or directory')
    parser.add_argument('-o', '--output', help='Output file or directory (defaults to input-{target_lang}.docx)')
    parser.add_argument('-s', '--source-lang', default='en', help='Source language code (default: en)')
    parser.add_argument('-t', '--target-lang', required=True, help='Target language code')
    
    # Advanced options
    parser.add_argument('--batch-size', type=int, default=50, help='Translation batch size')
    parser.add_argument('--concurrency', type=int, default=10, help='Max concurrent requests')
    parser.add_argument('--memory', default='translation_memory.txt', help='Translation memory file')
    parser.add_argument('--disable-memory', action='store_true', help='Disable translation memory')
    parser.add_argument('--verbose', '-v', action='store_true', help='Verbose output')
    
    args = parser.parse_args()
    
    # Configure logging
    log_level = logging.DEBUG if args.verbose else logging.INFO
    logging.basicConfig(
        level=log_level,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.StreamHandler(),
            logging.FileHandler('docx_translation.log')
        ]
    )
    logger = logging.getLogger('docx_translator')
    
    # Determine memory file
    memory_file = None if args.disable_memory else args.memory
    
    # Initialize translator
    translator = DocxTranslator(
        source_lang=args.source_lang,
        target_lang=args.target_lang,
        max_concurrent=args.concurrency,
        logger=logger,
        translation_memory_file=memory_file
    )
    
    try:
        # Check if target is RTL
        if args.target_lang in RTL_LANGUAGES:
            logger.info(f"Target language {args.target_lang} is RTL. RTL formatting will be applied.")
        
        # Different modes of operation
        if os.path.isdir(args.input):
            # Directory mode
            input_files = list(Path(args.input).glob('**/*.docx'))
            output_dir = args.output or os.path.join(os.path.dirname(args.input), f"translated_{args.target_lang}")
            
            logger.info(f"Found {len(input_files)} DOCX files in directory: {args.input}")
            stats = await translator.batch_translate_docx_files(
                input_files=[str(f) for f in input_files],
                output_dir=output_dir
            )
            
            # Print summary
            logger.info("\nTranslation Summary (Directory mode):")
            total_translated = sum(s.translated_chunks for s in stats.values())
            total_chunks = sum(s.total_text_chunks for s in stats.values())
            total_words = sum(s.translated_words for s in stats.values())
            total_chars = sum(s.character_count for s in stats.values())
            logger.info(f"Total: {total_translated}/{total_chunks} chunks, {total_words} words, {total_chars} characters translated")
            
        else:
            # Single file mode
            if not args.output:
                # Auto-generate output filename if not provided
                filename = os.path.basename(args.input)
                name, ext = os.path.splitext(filename)
                output_file = os.path.join(os.path.dirname(args.input), f"{name}_{args.target_lang}{ext}")
            else:
                output_file = args.output
            
            logger.info(f"Translating single DOCX file: {args.input} -> {output_file}")
            stats = await translator.translate_docx(
                input_file=args.input,
                output_file=output_file,
                batch_size=args.batch_size
            )
            
            # Print summary
            logger.info("\nTranslation Summary (Single file mode):")
            logger.info(f"Translated {stats.translated_chunks}/{stats.total_text_chunks} chunks")
            logger.info(f"Words: {stats.translated_words}/{stats.total_words}, Characters: {stats.character_count}")
            logger.info(f"Time taken: {stats.time_taken:.2f} seconds")
            
    except Exception as e:
        logger.error(f"Translation process failed: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    # Set up asyncio policies for Windows if needed
    if hasattr(asyncio, 'WindowsSelectorEventLoopPolicy') and hasattr(asyncio, 'set_event_loop_policy'):
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
    
    asyncio.run(main())
